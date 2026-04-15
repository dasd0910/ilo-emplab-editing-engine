"""
EMPLAB AI Editing and Proofreading Engine — main Streamlit app.

Phase 2 wired: deterministic ILO-style rule engine + Word tracked-changes writer.
Phase 3 (Gemini/Groq LLM pass), Phase 4 (IDML typesetting), and Phase 5 (PDF)
plug in after this.
"""

from __future__ import annotations

import hmac
import io
import os
from pathlib import Path

import streamlit as st
from docx import Document

from rules_engine import load_rules, find_edits, summarize_edits, resolve_overlaps
from docx_writer import apply_tracked_changes_with_comments
from llm_editor import llm_find_edits, load_api_keys, LLMError
from idml_writer import build_idml, get_template_path
from pdf_writer import build_pdf
from structural_rules import run_all_structural


# -----------------------------------------------------------------------------
# Page configuration
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="EMPLAB Editing Engine",
    page_icon="📝",
    layout="centered",
)


# -----------------------------------------------------------------------------
# Password gate — keeps the deployed app private to invited ILO colleagues.
# Password is read from Streamlit Cloud secrets (`app_password`) when deployed,
# or from the APP_PASSWORD env var when running locally. If no password is set,
# the gate is skipped (useful for local dev).
# -----------------------------------------------------------------------------
def _get_expected_password() -> str:
    # Streamlit Cloud: from st.secrets
    try:
        v = st.secrets.get("app_password")
        if v:
            return str(v)
    except Exception:
        pass
    # Local dev: from env / .env
    return os.environ.get("APP_PASSWORD", "")


def _password_gate() -> bool:
    """Return True if the user is authenticated (or no password required)."""
    expected = _get_expected_password()
    if not expected:  # open access when no password configured
        return True
    if st.session_state.get("auth_ok"):
        return True

    st.title("🔒 EMPLAB Editing Engine")
    st.caption("This tool is restricted to invited ILO colleagues.")
    pw = st.text_input("Password", type="password", key="pw_input")
    if st.button("Log in", type="primary") or pw:
        if pw and hmac.compare_digest(pw, expected):
            st.session_state["auth_ok"] = True
            st.rerun()
        elif pw:
            st.error("Incorrect password.")
    return False


if not _password_gate():
    st.stop()

RULES_PATH = Path(__file__).parent / "ilo_rules.yaml"
TEMPLATES_DIR = Path(__file__).parent / "templates"   # bundled .idml files
APP_ASSETS_DIR = Path(__file__).parent / "assets"     # HTML/CSS/fonts/logo for PDF


# Cache rules so we don't reload the YAML on every rerun
@st.cache_resource
def get_rules():
    return load_rules(RULES_PATH)


def _apply_plain(doc: Document, edits) -> None:
    """Apply auto-severity edits as plain text replacements (no tracked changes).
    Used to produce a clean Document for IDML export.
    """
    by_para: dict = {}
    for e in edits:
        if e.severity == "auto":
            by_para.setdefault(e.paragraph_index, []).append(e)
    for p_idx, para_edits in by_para.items():
        if p_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[p_idx]
        text = para.text
        if not text:
            continue
        # Apply edits right-to-left so earlier offsets remain valid
        para_edits.sort(key=lambda e: e.start, reverse=True)
        new_text = text
        for e in para_edits:
            new_text = new_text[:e.start] + e.replacement + new_text[e.end:]
        # Wipe existing runs and write the edited text into the first run
        for r in list(para.runs):
            r.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


# -----------------------------------------------------------------------------
# Header
# -----------------------------------------------------------------------------
st.title("📝 EMPLAB AI Editing & Proofreading Engine")
st.caption(
    "Upload an ILO draft. Pick a brief type. Get back a copy-edited Word file "
    "with tracked changes, plus (coming soon) a typeset InDesign file and "
    "print-ready PDF."
)
st.divider()


# -----------------------------------------------------------------------------
# Step 1 — upload
# -----------------------------------------------------------------------------
st.subheader("Step 1 — Upload your draft")
uploaded_file = st.file_uploader(
    "Drop a .docx file here (or click Browse)",
    type=["docx"],
    help="The draft publication you want to have copy-edited and typeset.",
)


# -----------------------------------------------------------------------------
# Step 2 — brief type
# -----------------------------------------------------------------------------
st.subheader("Step 2 — Select the brief type")
brief_type = st.selectbox(
    "Which ILO template should this publication use?",
    options=["Generic Brief", "Policy Brief", "Fact Sheet"],
    index=0,
    help="Determines the InDesign template used for typesetting (Phase 4).",
)


# -----------------------------------------------------------------------------
# Step 3 — options
# -----------------------------------------------------------------------------
st.subheader("Step 3 — Options")
col1, col2 = st.columns(2)
with col1:
    include_suggestions = st.checkbox(
        "Include low-confidence suggestions",
        value=False,
        help="If on, flags context-dependent edits (e.g., country expansions) "
             "for reviewer decision, without auto-applying them.",
    )
    use_llm = st.checkbox(
        "Run LLM grammar/flow pass (Phase 3)",
        value=True,
        help="Gemini (primary) / Groq (fallback). Catches grammar, clarity, "
             "tone, and awkward phrasing that rules can't. Uses your free-tier "
             "API quota. Takes 30–120 seconds depending on document length.",
    )
with col2:
    editor_author = st.text_input(
        "Tracked-change author name",
        value="ILO Editing Engine",
        help="Shown as the author of every tracked change in Word.",
    )


# -----------------------------------------------------------------------------
# Step 4 — run
# -----------------------------------------------------------------------------
st.subheader("Step 4 — Run the editor")
run_clicked = st.button("▶ Run editing pipeline", type="primary", use_container_width=True)

if run_clicked:
    if uploaded_file is None:
        st.error("Please upload a .docx file first.")
        st.stop()

    with st.spinner("Loading rules and reading document..."):
        rules = get_rules()
        raw_bytes = uploaded_file.getvalue()
        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text for p in doc.paragraphs]

    with st.spinner("Finding style-rule matches..."):
        rule_edits = find_edits(paragraphs, rules, include_suggestions=include_suggestions)

    with st.spinner("Applying structural rules (acronyms, headings, articles)..."):
        struct_edits = run_all_structural(doc, paragraphs)
        if not include_suggestions:
            # Structural rules are 'suggest' severity; only surface if user opted in
            struct_edits = []
        rule_edits = rule_edits + struct_edits

    # Phase 3 — LLM grammar/flow pass
    llm_edits = []
    if use_llm:
        keys = load_api_keys()
        if not any(keys.values()):
            st.warning(
                "No API keys found in `.env`. Skipping LLM pass. Set any of: "
                "`GEMINI_API_KEY`, `GROQ_API_KEY`, `CEREBRAS_API_KEY`, "
                "`MISTRAL_API_KEY` in your `.env` file to enable grammar/flow editing."
            )
        else:
            progress_bar = st.progress(0, text="LLM pass — starting…")
            def _cb(cur, total):
                progress_bar.progress(min(cur / max(total, 1), 1.0),
                                      text=f"LLM pass — chunk {cur} of {total}")
            try:
                llm_edits, llm_diag = llm_find_edits(
                    paragraphs,
                    gemini_key=keys["gemini"],
                    groq_key=keys["groq"],
                    cerebras_key=keys["cerebras"],
                    mistral_key=keys["mistral"],
                    progress_callback=_cb,
                )
                progress_bar.empty()
                # Surface LLM health so the user sees if quota ran out
                health = (
                    f"Chunks: {llm_diag['chunks_total']} total · "
                    f"Gemini: {llm_diag['chunks_ok_gemini']} · "
                    f"Groq: {llm_diag['chunks_ok_groq']} · "
                    f"Cerebras: {llm_diag.get('chunks_ok_cerebras', 0)} · "
                    f"Mistral: {llm_diag.get('chunks_ok_mistral', 0)} · "
                    f"Failed: {llm_diag['chunks_failed']}"
                )
                if llm_diag["chunks_failed"] > 0:
                    st.warning(f"LLM pass completed with some failures. {health}. "
                               f"Last error: `{llm_diag.get('last_error')}`")
                else:
                    st.info(f"LLM pass health: {health}")
            except LLMError as e:
                progress_bar.empty()
                st.error(f"LLM pass failed: {e}. Continuing with rule-engine edits only.")

    edits = resolve_overlaps(rule_edits + llm_edits)
    stats = summarize_edits(edits)

    with st.spinner("Writing tracked changes and comments..."):
        doc_to_edit = Document(io.BytesIO(raw_bytes))
        docx_bytes, result = apply_tracked_changes_with_comments(
            doc_to_edit, edits, author=editor_author, initials="IEE",
        )

    # Build one clean Document once — shared by IDML and PDF generators
    clean_doc_for_export = Document(io.BytesIO(raw_bytes))
    _apply_plain(clean_doc_for_export, edits)

    with st.spinner("Generating typeset IDML..."):
        idml_bytes = None
        idml_error = None
        try:
            template_path = get_template_path(brief_type, TEMPLATES_DIR)
            if not template_path.exists():
                idml_error = f"Bundled template missing: `{template_path.name}`."
            else:
                tmp_buf = io.BytesIO()
                clean_doc_for_export.save(tmp_buf); tmp_buf.seek(0)
                idml_bytes = build_idml(template_path, Document(tmp_buf), brief_type)
        except Exception as e:
            idml_error = f"{type(e).__name__}: {e}"

    with st.spinner("Generating print-ready PDF..."):
        pdf_bytes = None
        pdf_error = None
        try:
            tmp_buf = io.BytesIO()
            clean_doc_for_export.save(tmp_buf); tmp_buf.seek(0)
            pdf_bytes = build_pdf(
                Document(tmp_buf),
                brief_type=brief_type,
                assets_dir=APP_ASSETS_DIR,
                workspace_dir=APP_ASSETS_DIR,
            )
        except Exception as e:
            pdf_error = f"{type(e).__name__}: {e}"

    # Persist all results in session_state so they survive download-click reruns
    st.session_state["results"] = {
        "filename": uploaded_file.name,
        "brief_type": brief_type,
        "stats": stats,
        "result": result,
        "llm_edit_count": len(llm_edits),
        "edits_preview": [
            {"rule_id": e.rule_id, "manual_ref": e.manual_ref, "severity": e.severity,
             "original": e.original, "replacement": e.replacement,
             "description": e.description}
            for e in edits[:15]
        ],
        "total_edits": len(edits),
        "docx_bytes": docx_bytes,
        "idml_bytes": idml_bytes,
        "idml_error": idml_error,
        "pdf_bytes": pdf_bytes,
        "pdf_error": pdf_error,
    }


# -----------------------------------------------------------------------------
# Render results — driven by session_state so download clicks don't erase them
# -----------------------------------------------------------------------------
if "results" in st.session_state:
    r = st.session_state["results"]
    stats = r["stats"]
    result = r["result"]

    st.divider()
    st.success(
        f"✅ Done — applied {result['applied']} auto-edits "
        f"({result['suggested_skipped']} suggestions flagged but not applied). "
        f"LLM pass contributed {r['llm_edit_count']} edits."
    )

    c1, c2, c3 = st.columns(3)
    c1.metric("Edits found", stats["total"])
    c2.metric("Auto-applied", stats["auto"])
    c3.metric("Flagged for review", stats["suggest"])

    st.write(f"**File:** `{r['filename']}`  |  **Brief type:** {r['brief_type']}")

    if stats["by_category"]:
        st.write("**Edits by category:**")
        for cat, n in sorted(stats["by_category"].items(), key=lambda x: -x[1]):
            st.write(f"- `{cat}`: {n}")

    # Download buttons — persist because session_state holds the bytes
    stem = Path(r["filename"]).stem
    suffix = r["brief_type"].replace(" ", "_").lower()
    colA, colB, colC = st.columns(3)
    with colA:
        st.download_button(
            label="⬇ Edited .docx",
            data=r["docx_bytes"],
            file_name=f"{stem}_edited.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_docx",
            help="Word file with tracked changes + comments citing ILO style rules.",
        )
    with colB:
        if r["idml_bytes"] is not None:
            st.download_button(
                label="⬇ Typeset .idml",
                data=r["idml_bytes"],
                file_name=f"{stem}_{suffix}.idml",
                mime="application/vnd.adobe.indesign-idml-package",
                use_container_width=True,
                key="dl_idml",
                help="Opens in InDesign pre-populated with edited content.",
            )
        elif r["idml_error"]:
            st.error(f"IDML failed: {r['idml_error']}")
    with colC:
        if r["pdf_bytes"] is not None:
            st.download_button(
                label="⬇ Print-ready .pdf",
                data=r["pdf_bytes"],
                file_name=f"{stem}_{suffix}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="dl_pdf",
                help="ILO-branded A4 PDF for review and sharing.",
            )
        elif r["pdf_error"]:
            st.error(f"PDF failed: {r['pdf_error']}")

    # Preview
    if r["edits_preview"]:
        with st.expander(
            f"Preview: first {len(r['edits_preview'])} of {r['total_edits']} edits"
        ):
            for e in r["edits_preview"]:
                arrow = "→" if e["severity"] == "auto" else "⚑"
                st.write(
                    f"**{e['rule_id']}** (§{e['manual_ref']}, {e['severity']})  \n"
                    f"`{e['original']}` {arrow} `{e['replacement']}`  \n"
                    f"_{e['description']}_"
                )

    if st.button("🗑 Clear results", use_container_width=False):
        del st.session_state["results"]
        st.rerun()

    st.info(
        "✅ All three outputs ready. The **.docx** is for reviewer sign-off "
        "(tracked changes + rule-citing comments); the **.idml** goes to a "
        "designer for print-perfect layout; the **.pdf** is a shareable "
        "ILO-branded preview of the cleaned content."
    )


# -----------------------------------------------------------------------------
# Footer
# -----------------------------------------------------------------------------
st.divider()
st.caption("Built for the International Labour Organization · v0.5 (Phase 5)")
